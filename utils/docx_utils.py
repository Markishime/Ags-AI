"""
DOCX generation utilities for creating survey forms and documents
"""
import io
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_oil_palm_survey_docx():
    """
    Create a DOCX survey form for oil palm farmers based on the provided survey template.
    Returns the DOCX file as bytes.
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Title
    title = doc.add_heading('Oil Palm Farmer Survey: Soil and Leaf Testing Challenges', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Introduction paragraph
    intro = doc.add_paragraph()
    intro.add_run('This survey is intended for smallholder oil palm farmers in Malaysia. Its purpose is to understand the key challenges and bottlenecks they face in improving soil health. By identifying these barriers, we aim to reduce the quantities of chemical fertilizers used, which can help farmers lower their production costs. This reduction also brings environmental benefits by minimizing negative impacts on ecosystems and natural resources. Ultimately, the goal is to support more sustainable farming practices and improve farmer income through the application of AI-enabled agronomic tools.')
    
    # Section 1: Farmer Information
    doc.add_heading('Section 1: Farmer Information', level=1)
    
    # Question 1
    p1 = doc.add_paragraph()
    p1.add_run('1. Full Name: ').bold = True
    p1.add_run('_' * 50)
    
    # Question 2
    doc.add_paragraph('2. Farm Location').bold = True
    p2a = doc.add_paragraph()
    p2a.add_run('- State: ').bold = True
    p2a.add_run('_' * 30)
    
    p2b = doc.add_paragraph()
    p2b.add_run('- District: ').bold = True
    p2b.add_run('_' * 30)
    
    p2c = doc.add_paragraph()
    p2c.add_run('- Village/Area: ').bold = True
    p2c.add_run('_' * 30)
    
    # Question 3
    p3 = doc.add_paragraph()
    p3.add_run('3. How many hectares of oil palm do you manage? ').bold = True
    p3.add_run('_' * 10 + ' ha')
    
    # Section 2: Soil and Leaf Testing Practices
    doc.add_heading('Section 2: Soil and Leaf Testing Practices', level=1)
    
    # Question 4
    p4 = doc.add_paragraph()
    p4.add_run('4. Do you currently do soil testing on your farm?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Other: ' + '_' * 30)
    
    # Question 5
    p5 = doc.add_paragraph()
    p5.add_run('5. Do you currently do leaf (plant tissue) testing?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Other: ' + '_' * 30)
    
    # Question 6
    p6 = doc.add_paragraph()
    p6.add_run('6. How often do you carry out these tests?').bold = True
    doc.add_paragraph('☐ Once a year ☐ Every 2–3 years ☐ Only when there is a problem ☐ Never')
    doc.add_paragraph('☐ Other: ' + '_' * 30)
    
    # Question 7
    p7 = doc.add_paragraph()
    p7.add_run('7. Can you interpret the test results easily?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Partially   ☐ Other: ' + '_' * 20)
    
    # Question 8
    p8 = doc.add_paragraph()
    p8.add_run('8. Do you create a fertilization plan based on the test results?').bold = True
    doc.add_paragraph('☐ Yes ☐ No ☐ Someone else does it for me ☐ Other: ' + '_' * 20)
    
    # Question 9
    p9 = doc.add_paragraph()
    p9.add_run('9. (If Yes to Q8) Is it easy for you to follow the fertilization plan?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Somewhat   ☐ Other: ' + '_' * 20)
    
    # Section 3: Support and Advice
    doc.add_heading('Section 3: Support and Advice', level=1)
    
    # Question 10
    p10 = doc.add_paragraph()
    p10.add_run('10. Have you ever received written guidelines or plans from agronomists or officers?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Cannot remember   ☐ Other: ' + '_' * 20)
    
    # Question 11
    p11 = doc.add_paragraph()
    p11.add_run('11. How satisfied are you with the support you receive from agronomists or extension officers?').bold = True
    doc.add_paragraph('☐ Very satisfied   ☐ Somewhat satisfied   ☐ Not satisfied   ☐ Never received any support   ☐ Other: ' + '_' * 20)
    
    # Question 12
    p12 = doc.add_paragraph()
    p12.add_run('12. Do you ask an agronomist or extension officer for help with test results?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Other: ' + '_' * 30)
    
    # Question 13
    p13 = doc.add_paragraph()
    p13.add_run('13. (If Yes) Do they give you a fertilization plan?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Sometimes   ☐ Other: ' + '_' * 20)
    
    # Question 14
    p14 = doc.add_paragraph()
    p14.add_run('14. (If Yes) Do they help you understand your return on investment (how much extra income your investments will generate)?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Other: ' + '_' * 20)
    
    # Question 15
    p15 = doc.add_paragraph()
    p15.add_run('15. Do they provide:').bold = True
    doc.add_paragraph('☐ Short-term improvement plans ☐ Long-term improvement plans ☐ Both ☐ None')
    doc.add_paragraph('☐ Other: ' + '_' * 30)
    
    # Section 4: Your Needs and Interest
    doc.add_heading('Section 4: Your Needs and Interest', level=1)
    
    # Question 16
    p16 = doc.add_paragraph()
    p16.add_run('16. Do you currently pay an agronomist for advice or support?').bold = True
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Other: ' + '_' * 30)
    
    # Sub-question for 16
    doc.add_paragraph('(If Yes) How much do you usually pay per visit or per year?')
    p16a = doc.add_paragraph()
    p16a.add_run('Amount: RM ').bold = True
    p16a.add_run('_' * 15 + '   Frequency: ' + '_' * 20)
    
    # Question 17
    p17 = doc.add_paragraph()
    p17.add_run('17. Would you be willing to pay a small fee for a service that gives you:').bold = True
    doc.add_paragraph('- Clear interpretation of soil and leaf test results')
    doc.add_paragraph('- Fertilizer and soil improvement strategies (short + long term)')
    doc.add_paragraph('- Forecast of expected economic returns based on your investment')
    doc.add_paragraph('☐ Yes   ☐ No      ☐ Other: ' + '_' * 20)
    
    # Question 18
    p18 = doc.add_paragraph()
    p18.add_run('18. What other agronomy problems do you need help with?').bold = True
    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)
    doc.add_paragraph('_' * 50)
    
    # Question 19
    p19 = doc.add_paragraph()
    p19.add_run('19. Would you like to join a free trial of a new AI-based tool that:').bold = True
    doc.add_paragraph()
    doc.add_paragraph('- Clearly interprets your soil and leaf test results')
    doc.add_paragraph('- Provides targeted short- and long-term improvement strategies')
    doc.add_paragraph('- Estimates how your profits could increase based on the actions you take')
    doc.add_paragraph('☐ Yes   ☐ No   ☐ Other: ' + '_' * 30)
    
    # Contact details section
    doc.add_paragraph()
    doc.add_paragraph('If you answered YES please provide your contact details:').bold = True
    p20a = doc.add_paragraph()
    p20a.add_run('Email Address: ').bold = True
    p20a.add_run('_' * 40)
    
    p20b = doc.add_paragraph()
    p20b.add_run('WhatsApp Number: ').bold = True
    p20b.add_run('_' * 30)
    
    # Privacy notice
    doc.add_paragraph()
    privacy = doc.add_paragraph('We will not share this information with anyone. It will only be used to send you with instructions on how to use the AI-based tool for the trial.')
    privacy.italic = True
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer.getvalue()


def create_survey_download_button():
    """
    Create a download button for the oil palm farmer survey DOCX file.
    Returns the button widget and file data.
    """
    try:
        docx_data = create_oil_palm_survey_docx()
        return docx_data
    except Exception as e:
        raise Exception(f"Error creating survey DOCX: {str(e)}") from e
